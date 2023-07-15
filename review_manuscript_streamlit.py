from langchain.document_loaders import PyPDFLoader #pip install pypdf==3.12.1
from IPython.display import display, Markdown
import openai
import os
# from dotenv import load_dotenv, find_dotenv
# _ = load_dotenv(find_dotenv()) # read local .env file with OpenAI key
# openai.api_key = openai_api_key=os.environ['OPENAI_API_KEY']
from dotenv import load_dotenv
load_dotenv()
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', 'YourAPIKeyIfNotSet')

from langchain import OpenAI
from langchain.text_splitter import TokenTextSplitter
import tiktoken
import warnings
warnings.filterwarnings("ignore")
import time

import streamlit as st
import pdfplumber
import docx
import io

# A function that will be called only if the environment's openai_api_key isn't set
def get_openai_api_key():
    input_text = st.text_input(label="OpenAI API Key",  placeholder="Ex: sk-2twmA8tfCb8un4...", key="openai_api_key_input")
    return input_text

# Start Of Streamlit page
st.set_page_config(page_title="Paper review assistant", page_icon=":robot:")

# Start Top Information
st.header("LLM Assisted Paper review")

col1, col2 = st.columns(2)

with col1:
    st.markdown("Have a manuscript to review? Takes a lot of time? This tool is meant to help you generate \
                questions for the manuscript that can help improving it. There is a limit of 16k tokens (with response). Ensure that your pdf doesn't exceed 15k tokens. \
                \n\nThis tool is made  by [Nikos Sourlos](www.linkedin.com/in/nsourlos). \n\n View Source Code on [Github](https://github.com/nsourlos/review_manuscript_streamlit/blob/main/review_manuscript_streamlit.py)")

with col2:
    # st.image(image='paper_review.jpg', width=300, caption='https://www.nature.com/articles/d41586-018-06991-0')
    # st.markdown("![Alt Text](https://github.com/nsourlos/review_manuscript_streamlit/blob/main/paper_review.gif)")
    #   st.image(
    #         "https://github.com/nsourlos/review_manuscript_streamlit/blob/main/paper_review.gif", # I prefer to load the GIFs using GIPHY
    #         width=400, # The actual size of most gifs on GIPHY are really small, and using the column-width parameter would make it weirdly big. So I would suggest adjusting the width manually!
    #     )
    import base64

    file_ = open("C:/Users/soyrl/Desktop/review_manuscript_streamlit/paper_review.gif", "rb")
    contents = file_.read()
    data_url = base64.b64encode(contents).decode("utf-8")
    file_.close()

    st.markdown(
        f'<img src="data:image/gif;base64,{data_url}" alt="paper gif">',
        unsafe_allow_html=True,
    )

     #https://discuss.streamlit.io/t/how-to-show-local-gif-image/3408/3
# End Top Information

st.markdown("## :muscle: Upload PDF document")
	# :older_man:
# Output type selection by the user
# output_type = st.radio(
#     "Output Type:",
#     ('Interview Questions', '1-Page Summary'))

uploaded_file = st.file_uploader("Choose a pdf file", type="pdf") #https://discuss.streamlit.io/t/how-to-upload-a-pdf-file-in-streamlit/2428/2
# if uploaded_file.name.endswith(".pdf")==0:
#     st.write("Only accepts PDF files. Please select another file")
# else:

# # Collect information about the person you want to research
OPENAI_API_KEY = st.text_input(label="OpenAI API Key (or set it as .env variable)",  placeholder="Ex: sk-2twmA8tfCb8un4...", key="YourAPIKeyIfNotSet")
# OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', 'YourAPIKeyIfNotSet')
# st.write("OpenAI API Key:", OPENAI_API_KEY)

# download_option = st.selectbox("Download questions as docx?",('Yes', 'No'),index=1) #https://stackoverflow.com/questions/65026852/set-default-value-for-selectbox
# twitter_handle = st.text_input(label="Twitter Username",  placeholder="@eladgil", key="twitter_user_input")
# youtube_videos = st.text_input(label="YouTube URLs (Use , to seperate videos)",  placeholder="Ex: https://www.youtube.com/watch?v=c_hO_fjmMnk, https://www.youtube.com/watch?v=c_hO_fjmMnk", key="youtube_user_input")
# webpages = st.text_input(label="Web Page URLs (Use , to seperate urls. Must include https://)",  placeholder="https://eladgil.com/", key="webpage_user_input")

button_ind = st.button("*Generate Output*", type='secondary', help="Click to generate review questions")
if button_ind:
    if uploaded_file is None:
        st.warning('Please provide a PDF file', icon="⚠️")
        st.stop()
    
    # if uploaded_file.name.endswith('.pdf')==0:
    #     #More icons in https://streamlit-emoji-shortcodes-streamlit-app-gwckff.streamlit.app/
    #     st.warning('Only accepts PDF files. Please select another file', icon="🚨")
    #     st.stop()

    if not OPENAI_API_KEY:
        st.warning('Please insert OpenAI API Key. Instructions [here](https://help.openai.com/en/articles/4936850-where-do-i-find-my-secret-api-key)', icon="⚠️")
        st.stop()

    if OPENAI_API_KEY == 'YourAPIKeyIfNotSet':
        st.warning('Please fill in OpenAI API Key. Instructions [here](https://help.openai.com/en/articles/4936850-where-do-i-find-my-secret-api-key)', icon="⚠️")
        st.stop()
        # If the openai key isn't set in the env, put a text box out there
        # while OPENAI_API_KEY == 'YourAPIKeyIfNotSet':
            # st.session_state["data"] = {OPENAI_API_KEY: None}
            # result=st.session_state["data"].get(OPENAI_API_KEY, None)
            # st.write("result", result)
            # while result is None:
                # OPENAI_API_KEY = get_openai_api_key()
                # time.sleep(5)
                # result=st.session_state["data"].get(OPENAI_API_KEY, None)
                # if result is not None:
                #     st.write("key given", result)
            # st.experimental_rerun()





    #Information on how to load uploaded file in https://docs.streamlit.io/library/api-reference/widgets/st.file_uploader
    #For pdfs https://discuss.streamlit.io/t/how-to-upload-a-pdf-file-in-streamlit/2428
    st.write("Loading PDF...")
    manuscript_path=uploaded_file#.getvalue()#'paper.pdf'
    # print(manuscript_path)
    # print()
    # st.write("Loaded PDF dir", os.getcwd()+'/'+manuscript_path.name) # /app/review_manuscript_streamlit/paper.pdf
    # st.write("Loaded PDF", manuscript_path) #UploadedFile(id=9, name='paper.pdf', type='application/pdf', size=3197009)
    # st.write("Loaded PDF name not used", uploaded_file.name) #paper.pdf
    # st.write("Loading PDF get value", uploaded_file.getvalue())
    # st.write("Loading PDF read", uploaded_file.read()) #TypeError: A bytes-like object is required, not 'str'

    # st.write("Loading PDF get value decode", uploaded_file.getvalue().decode("utf-8")) 
    #UnicodeDecodeError: 'utf-8' codec can't decode byte 0xbf in position 10: invalid start byte


    # Load PDF
    # try:
    #     loaders = [
    #         PyPDFLoader(manuscript_path), #ypeError: stat: path should be string, bytes, os.PathLike or integer, not UploadedFile
    #     ]
    # except:
    #     import traceback
    #     st.write(traceback.format_exc())

        # try:
        #     buffer = io.BytesIO()
        #     # write binary content of file to buffer
        #     with open(manuscript_path, mode='rb') as file: 
        #         buffer.write(file.read())

        #     st.write("Loaded PDF buffer", buffer)
        # except:
        #     st.write(traceback.format_exc())

    pdf_file=pdfplumber.open(manuscript_path)#no attribute load
        # st.write("Loaded PDF dir inside", pdf_file) #info on pdfplumber object
        # st.write("Loaded PDF dir pages", pdf_file.pages) #with len get num of pages
        # st.write("Loaded PDF dir 1 page text", pdf_file.pages[0].extract_text(x_tolerance=1))

    docs = [] #https://github.com/jsvine/pdfplumber/issues/147
    for page in pdf_file.pages: #Add all documents to one
        docs.append(page.extract_text(x_tolerance=1)) #https://github.com/jsvine/pdfplumber/issues/334
    pdf_file.close()
    # st.write("Loaded PDF docs", docs)
    # paper=[docs[i].page_content for i in range(len(docs))] #Get only document content
    paper=''.join(docs)
    # st.write("final PDF", paper)

    # docs = []
    # for loader in loaders: #Add all documents to one
    #     docs.extend(loader.load())
    # paper=[docs[i].page_content for i in range(len(docs))] #Get only document content
    # paper=''.join(paper)

    #Calculate token usage and price for CV prompt - Same as sent to OpenAI but free
    st.write("Calculating price...")
    text_splitter = TokenTextSplitter(chunk_size=1, chunk_overlap=0,model_name='gpt-3.5-turbo-16k')
    ind_tokens=text_splitter.split_text(paper)
    # print("Total tokens:",len(ind_tokens))
    # print("Price for them:",round(len(ind_tokens)*0.003/1000,4),"$")
    st.write("Total tokens:",len(ind_tokens))
    st.write("Price for them:",round(len(ind_tokens)*0.003/1000,4),"$")
    st.write("Loading LLM output...")

    review_prompt='Below is a manuscript of a scientific publication. Act as a reviewer for the manuscript and provide at least 10 points for improvement. \
        These points should provide be related to the content of the manuscript. The manuscript is: '
    llm=OpenAI(openai_api_key=OPENAI_API_KEY,temperature=0,model_name='gpt-3.5-turbo-16k') #Initialize LLM - 16k context length to fit the paper
    questions_final=llm.predict(review_prompt+paper) #Predict response using LLM 
    display(Markdown(questions_final))
    st.markdown(f"#### Review Questions:")
    st.write(questions_final)

    document = docx.Document() #Create word document
    document.add_heading('Review Questions', level=1) #Add title
    document.add_paragraph(questions_final) #Add text
    # final_file=document.save('review_questions.docx') #Save document
    # st.write('doc',document) #aatributes pritned

#     st.download_button(
#         label="Download questions as .docx",
#         data=document,
#         file_name='review_questions.docx',
#         # mime='text/csv',
# )

    #https://discuss.streamlit.io/t/downloading-a-ms-word-document/28850/3
    bio = io.BytesIO()
    document.save(bio)
    if document:
        st.download_button(
            label="Click here to download as 'docx'",
            data=bio.getvalue(),
            file_name="review_questions.docx",
            mime="docx"
        )
    # if download_option=='Yes':
    #     document = docx.Document() #Create word document
    #     document.add_heading('Review Questions', level=1) #Add title
    #     document.add_paragraph(questions_final) #Add text
    #     final_file=document.save('review_questions.docx') #Save document
        

#Notes: Streamlit gives 1CPU, 1GB of RAM and 1GB of disk space (https://discuss.streamlit.io/t/problem-on-resources-limit/12605)
#Streamlit implementation adapted from https://github.com/gkamradt/llm-interview-research-assistant/blob/main/main.py
#Deployment based on https://docs.streamlit.io/streamlit-community-cloud/get-started/deploy-an-app